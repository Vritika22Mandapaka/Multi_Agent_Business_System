import os
import sys
import json
import tempfile
from io import BytesIO
from datetime import datetime

import fitz
import streamlit as st
from docx import Document

from app.graph import build_graph
from evals.consistency_check import run_consistency_check
from evals.rubric_eval import run_rubric_evaluation
from rag.extract_state import extract_state
from rag.domain_detector import detect_business_domain


st.set_page_config(
    page_title="AI Multi-Agent Business Decision System",
    page_icon="🤖",
    layout="wide",
)


st.markdown(
    """
    <style>
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }

    .hero-box {
        background: linear-gradient(135deg, #0f172a 0%, #1e3a8a 55%, #2563eb 100%);
        padding: 2rem;
        border-radius: 18px;
        color: white;
        margin-bottom: 1.5rem;
        box-shadow: 0px 8px 24px rgba(15, 23, 42, 0.18);
    }

    .hero-title {
        font-size: 2.2rem;
        font-weight: 800;
        margin-bottom: 0.35rem;
    }

    .hero-subtitle {
        font-size: 1rem;
        opacity: 0.92;
        line-height: 1.6;
    }

    .metric-card {
        background-color: white;
        padding: 1.1rem;
        border-radius: 16px;
        border: 1px solid #e5e7eb;
        box-shadow: 0px 4px 14px rgba(15, 23, 42, 0.06);
        height: 100%;
    }

    .section-card {
        background-color: white;
        padding: 1.25rem;
        border-radius: 16px;
        border: 1px solid #e5e7eb;
        box-shadow: 0px 4px 14px rgba(15, 23, 42, 0.06);
        margin-bottom: 1rem;
    }

    .small-muted {
        color: #64748b;
        font-size: 0.9rem;
    }

    .success-pill {
        display: inline-block;
        background-color: #dcfce7;
        color: #166534;
        padding: 0.35rem 0.7rem;
        border-radius: 999px;
        font-weight: 700;
        font-size: 0.85rem;
        margin-right: 0.4rem;
        margin-top: 0.5rem;
    }

    .warn-pill {
        display: inline-block;
        background-color: #fef3c7;
        color: #92400e;
        padding: 0.35rem 0.7rem;
        border-radius: 999px;
        font-weight: 700;
        font-size: 0.85rem;
        margin-right: 0.4rem;
        margin-top: 0.5rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


def extract_text_from_pdf(uploaded_file):
    text = ""

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
        temp_file.write(uploaded_file.read())
        temp_path = temp_file.name

    try:
        doc = fitz.open(temp_path)
        try:
            for page in doc:
                text += page.get_text()
        finally:
            doc.close()
    finally:
        os.remove(temp_path)

    return text.strip()


def run_multi_agent_system(business_text):
    app = build_graph()

    target_state = extract_state(business_text)
    business_domain = detect_business_domain(business_text)

    initial_state = {
        "business_idea": business_text,
        "business_domain": business_domain,

        "research_output": None,
        "research_eval": None,

        "tech_output": None,
        "tech_eval": None,

        "finance_output": None,

        "retry_output": None,

        "final_report": None,

        "target_state": target_state,
        "regulatory_context": None,
    }

    return app.invoke(initial_state)


def safe_get_score(eval_dict):
    if not eval_dict:
        return "N/A"

    score = eval_dict.get("rubric_score", eval_dict.get("score"))
    total = eval_dict.get("max_score", eval_dict.get("total"))

    if score is None or total is None:
        return "N/A"

    return f"{score}/{total}"


def dict_to_pretty_text(data):
    if data is None:
        return "Not available"

    if isinstance(data, str):
        return data

    return json.dumps(data, indent=2, ensure_ascii=False)


def generate_docx_report(final_state, business_text):
    document = Document()

    document.add_heading("AI Multi-Agent Business Decision Report", 0)

    document.add_paragraph(
        f"Generated On: {datetime.now().strftime('%B %d, %Y %I:%M %p')}"
    )
    document.add_paragraph("Project: AI Multi-Agent Business Decision System")

    document.add_heading("Business Idea", level=1)
    document.add_paragraph(business_text)

    document.add_heading("Detected Context", level=1)
    document.add_paragraph(
        f"Detected Business Domain: {final_state.get('business_domain', 'general')}"
    )
    document.add_paragraph(
        f"Detected State: {final_state.get('target_state') or 'Not detected'}"
    )

    document.add_heading("Final Business Decision Report", level=1)
    document.add_paragraph(final_state["final_report"]["report"])

    document.add_heading("Research Agent Output", level=1)
    document.add_paragraph(
        dict_to_pretty_text(final_state.get("research_output", {}).get("analysis"))
    )

    document.add_heading("Technology Agent Output", level=1)
    document.add_paragraph(dict_to_pretty_text(final_state.get("tech_output")))

    document.add_heading("Finance Agent Output", level=1)
    document.add_paragraph(dict_to_pretty_text(final_state.get("finance_output")))

    if final_state.get("retry_output"):
        document.add_heading("Retry Agent Output", level=1)
        document.add_paragraph(dict_to_pretty_text(final_state.get("retry_output")))

    document.add_heading("Evaluation Results", level=1)

    document.add_paragraph("Research Evaluation:")
    document.add_paragraph(dict_to_pretty_text(final_state.get("research_eval")))

    document.add_paragraph("Technology Evaluation:")
    document.add_paragraph(dict_to_pretty_text(final_state.get("tech_eval")))

    if final_state.get("regulatory_context"):
        document.add_heading("Retrieved Regulatory Context", level=1)
        document.add_paragraph(final_state["regulatory_context"][:6000])

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer


def render_eval_card(title, eval_data):
    score_text = safe_get_score(eval_data)

    st.markdown(
        f"""
        <div class="metric-card">
            <div class="small-muted">{title}</div>
            <h3>{score_text}</h3>
        </div>
        """,
        unsafe_allow_html=True,
    )


st.markdown(
    """
    <div class="hero-box">
        <div class="hero-title">AI Multi-Agent Business Decision System</div>
        <div class="hero-subtitle">
            Evaluate business ideas using specialized AI agents for market research,
            technical feasibility, finance, regulatory retrieval, and final strategic synthesis.
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)


with st.sidebar:
    st.header("System Pipeline")

    st.markdown(
        """
        **1. Domain Detection**  
        Identifies business category

        **2. Research Agent**  
        Market RAG when domain matches + generic fallback

        **3. Technology Agent**  
        Architecture + implementation plan

        **4. Finance Agent**  
        Unit economics + ROI analysis

        **5. Regulatory RAG**  
        State-specific compliance retrieval

        **6. Synthesis Agent**  
        Final Go / No-Go decision
        """
    )

    st.divider()
    st.caption(f"Python executable: `{sys.executable}`")


st.markdown("## Submit a Business Idea")

tab1, tab2 = st.tabs(["Upload File", "Type Manually"])

business_text = ""

with tab1:
    uploaded_file = st.file_uploader(
        "Upload a PDF or TXT file",
        type=["pdf", "txt"],
        help="Upload a business idea document.",
    )

    if uploaded_file:
        if uploaded_file.name.endswith(".pdf"):
            business_text = extract_text_from_pdf(uploaded_file)
        else:
            business_text = uploaded_file.read().decode("utf-8").strip()

        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.subheader("Input Preview")
        st.text_area("Business Idea Text", business_text, height=180)
        st.markdown("</div>", unsafe_allow_html=True)

with tab2:
    typed_text = st.text_area(
        "Enter your business idea",
        height=180,
        placeholder=(
            "Example 1: I want to launch an AI-powered grocery delivery startup "
            "in New Jersey for college students.\n\n"
            "Example 2: I want to start a preschool program in New Jersey."
        ),
    )

    if typed_text.strip():
        business_text = typed_text.strip()


if st.button("Run Multi-Agent Analysis", type="primary", use_container_width=True):
    if not business_text:
        st.error("Please upload a file or type a business idea first.")
        st.stop()

    with st.spinner(
        "Running Domain Detection, Research, Tech, Finance, Regulatory RAG, and Synthesis Agents..."
    ):
        final_state = run_multi_agent_system(business_text)

    st.session_state["final_state"] = final_state
    st.session_state["business_text"] = business_text


if "final_state" in st.session_state:
    final_state = st.session_state["final_state"]
    business_text = st.session_state["business_text"]

    st.success("Analysis completed successfully.")

    detected_state = final_state.get("target_state")
    detected_domain = final_state.get("business_domain", "general")

    if detected_state:
        st.markdown(
            f'<span class="success-pill">Detected State: {detected_state}</span>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            '<span class="warn-pill">No supported state detected. Regulatory RAG skipped.</span>',
            unsafe_allow_html=True,
        )

    st.markdown(
        f'<span class="success-pill">Detected Domain: {detected_domain}</span>',
        unsafe_allow_html=True,
    )

    st.markdown("## Executive Summary")

    col1, col2, col3 = st.columns(3)

    with col1:
        render_eval_card("Research Evaluation", final_state.get("research_eval"))

    with col2:
        render_eval_card("Technology Evaluation", final_state.get("tech_eval"))

    with col3:
        regulatory_status = (
            "Retrieved"
            if final_state.get("regulatory_context")
            else "Not Available"
        )

        st.markdown(
            f"""
            <div class="metric-card">
                <div class="small-muted">Regulatory Context</div>
                <h3>{regulatory_status}</h3>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("## Final Business Decision Report")

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown(final_state["final_report"]["report"])
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("## Agent-Level Outputs")

    agent_tab1, agent_tab2, agent_tab3, agent_tab4, agent_tab5 = st.tabs(
        [
            "Research",
            "Technology",
            "Finance",
            "Regulatory Context",
            "Evaluation",
        ]
    )

    with agent_tab1:
        st.subheader("Research Agent")

        st.info(
            f"Detected Business Domain: {final_state.get('business_domain', 'general')}"
        )

        research_output = final_state.get("research_output", {})

        if research_output.get("rag_usage_note"):
            st.caption(research_output.get("rag_usage_note"))

        st.json(research_output.get("analysis"))

        with st.expander("Market RAG Context"):
            market_context = research_output.get(
                "market_rag_context",
                "Not available",
            )

            if market_context:
                st.text(market_context)
            else:
                st.info("Market RAG was skipped for this domain.")

        with st.expander("Web / Market Context Tool"):
            st.json(research_output.get("web_market_context", {}))

    with agent_tab2:
        st.subheader("Technology Agent")
        st.json(final_state.get("tech_output"))

    with agent_tab3:
        st.subheader("Finance Agent")
        st.json(final_state.get("finance_output"))

    with agent_tab4:
        st.subheader("Regulatory Context Summary")

        if final_state.get("regulatory_context"):
            regulatory_text = final_state["regulatory_context"]

            st.success("Regulatory RAG successfully retrieved official context.")

            preview = regulatory_text[:1800]

            st.markdown("### Key Retrieved Evidence Preview")
            st.text_area(
                "Top Retrieved Regulatory Context",
                preview,
                height=250,
            )

            with st.expander("View Full Raw Regulatory Retrieval"):
                st.text(regulatory_text)
        else:
            st.info("No regulatory context was retrieved.")

    with agent_tab5:
        st.subheader("Rubric-Based Evaluation")

        evaluation_results = run_rubric_evaluation(final_state)

        for agent, result in evaluation_results.items():
            st.write(f"**{agent}:** {result['score']} / {result['total']}")

            if result["missing"]:
                st.warning(f"Missing: {', '.join(result['missing'])}")

        st.divider()

        st.subheader("Consistency Check")

        if st.button("Run Consistency Check"):
            with st.spinner("Running second pass for consistency validation..."):
                second_run_state = run_multi_agent_system(business_text)

            consistency_result = run_consistency_check(
                final_state["final_report"]["report"],
                second_run_state["final_report"]["report"],
            )

            st.metric(
                label="Consistency Score",
                value=f"{consistency_result['consistency_score_percent']}%",
            )

    st.markdown("## Download Report")

    docx_file = generate_docx_report(final_state, business_text)

    st.download_button(
        label="Download Final Report as Word Document",
        data=docx_file,
        file_name="multi_agent_business_decision_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )